"""
画面仕様書 / 状態遷移図 / ユーザーストーリー / テストケース
Word (.docx) + Excel (.xlsx) 生成スクリプト
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from pathlib import Path

OUT_DIR = Path(__file__).parent

# ──────────────────────────────────────────────
# 共通ユーティリティ
# ──────────────────────────────────────────────

def hex_fill(h): return PatternFill("solid", fgColor=h)
def thin_border():
    s = Side(style="thin", color="BFBFBF")
    return Border(left=s, right=s, top=s, bottom=s)
def wrap(h="left", v="top"):
    return Alignment(horizontal=h, vertical=v, wrap_text=True)

def xl_header(ws, row, cols, bg="1F4E79"):
    for c in cols:
        cell = ws.cell(row=row, column=c)
        cell.fill = hex_fill(bg)
        cell.font = Font(name="メイリオ", bold=True, color="FFFFFF", size=9)
        cell.alignment = wrap("center", "center")
        cell.border = thin_border()

def xl_cell(ws, row, col, val, bold=False, bg=None, h="left", v="top", size=9):
    cell = ws.cell(row=row, column=col, value=val)
    cell.font = Font(name="メイリオ", bold=bold, size=size)
    cell.alignment = wrap(h, v)
    cell.border = thin_border()
    if bg: cell.fill = hex_fill(bg)
    return cell

def xl_title(ws, text, cols):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=cols)
    c = ws.cell(row=1, column=1, value=text)
    c.font = Font(name="メイリオ", bold=True, size=13, color="FFFFFF")
    c.fill = hex_fill("1F4E79")
    c.alignment = wrap("center", "center")
    ws.row_dimensions[1].height = 26

def set_w(ws, col, cm): ws.column_dimensions[get_column_letter(col)].width = cm * 4.5

# Word ユーティリティ
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color); shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def w_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if p.runs:
        r = p.runs[0]
        if level == 1: r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79); r.font.size = Pt(15)
        elif level == 2: r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6); r.font.size = Pt(12)
        elif level == 3: r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79); r.font.size = Pt(10)

def w_table(doc, headers, rows, widths, hdr_bg="1F4E79"):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i,(h,w) in enumerate(zip(headers,widths)):
        c = tbl.cell(0,i); c.text=h; c.width=Cm(w)
        r = c.paragraphs[0].runs[0]
        r.font.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, hdr_bg)
    stripe = ["FFFFFF","EEF3FA"]
    for ri,row in enumerate(rows):
        bg = stripe[ri%2]
        for ci,val in enumerate(row):
            c = tbl.cell(ri+1,ci); c.text=str(val); c.width=Cm(widths[ci])
            r=c.paragraphs[0].runs[0]; r.font.size=Pt(8)
            if ci==0: r.font.bold=True
            set_cell_bg(c, bg)
    return tbl

def w_para(doc, text, size=9, italic=False, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.italic = italic
    return p

# ──────────────────────────────────────────────
# データ定義
# ──────────────────────────────────────────────

SCREEN_SPECS = [
    {
        "id": "SCR-001", "name": "ログイン画面", "file": "index.html",
        "url": "/index.html（/）",
        "遷移元": "初回アクセス / ログアウト後 / PW変更完了後",
        "概要": "ユーザーIDとパスワードを入力して認証を行う画面。",
        "入力項目": [
            ("ユーザーID", "text", "1〜20文字 / 半角英数字", "必須 / 文字種 / 文字数", "blur & 送信時"),
            ("パスワード", "password（マスク）", "8〜32文字 / 半角英数記号", "必須 / 文字種 / 文字数", "blur & 送信時"),
        ],
        "ボタン": [("ログイン", "POST /api/auth/login を呼び出す")],
        "処理フロー": [
            "① フロントエンドバリデーション（エラーがあれば停止）",
            "② POST /api/auth/login",
            "　・ロック中 → #auth-error（locked スタイル）",
            "　・認証失敗 → #auth-error（残り試行回数付き）",
            "　・期限切れ（422）→ SCR-003（PW変更画面）へ切替",
            "　・管理者ログイン成功 → /admin.html へリダイレクト",
            "　・一般ユーザーログイン成功 → SCR-002（ホーム画面）へ切替",
        ],
        "エラー表示": [
            ("フィールドエラー", "各入力欄直下", "#e53e3e", "フォーカスアウト / 送信時"),
            ("認証エラー", "入力欄上部（共通）", "#c53030（薄赤背景）", "API 応答時"),
            ("ロックエラー", "入力欄上部（共通）", "#744210（薄黄背景）", "API 403 時"),
        ],
        "備考": "入力開始時（input イベント）に認証エラーメッセージを非表示にする。"
    },
    {
        "id": "SCR-002", "name": "ホーム画面", "file": "index.html",
        "url": "同一ページ（SPA）",
        "遷移元": "一般ユーザーのログイン成功時",
        "概要": "ログイン成功後に表示されるダッシュボード画面。",
        "入力項目": [],
        "ボタン": [("ログアウト", "POST /api/auth/logout → SCR-001 へ切替・フォームをクリア")],
        "処理フロー": [
            "① ヘッダーに「{userId} さん」を表示",
            "② ログイン成功メッセージ（日時付き）を表示",
            "③ 6機能のダッシュボードカードを表示（遷移なし）",
        ],
        "エラー表示": [],
        "備考": ""
    },
    {
        "id": "SCR-003", "name": "パスワード変更画面", "file": "index.html",
        "url": "同一ページ（SPA）",
        "遷移元": "PW期限切れユーザーのログイン成功時",
        "概要": "パスワード有効期限（90日）が切れた場合に強制表示される変更画面。",
        "入力項目": [
            ("新しいパスワード", "password（マスク）", "8〜32文字 / 半角英数記号", "必須 / 文字種 / 文字数", "送信時"),
            ("新しいパスワード（確認）", "password（マスク）", "—", "必須 / 新PWとの一致", "送信時"),
        ],
        "ボタン": [("パスワードを変更する", "POST /api/auth/change-password")],
        "処理フロー": [
            "① フロントエンドバリデーション（新PW + 確認PW）",
            "② POST /api/auth/change-password",
            "　・成功 → 成功メッセージ表示 → ボタン disabled → 2秒後 SCR-001 へ",
            "　・失敗 → フィールドエラー表示",
        ],
        "エラー表示": [
            ("新PW フィールドエラー", "#err-new-password", "#e53e3e", "送信時"),
            ("確認PW フィールドエラー", "#err-confirm-password", "#e53e3e", "送信時"),
        ],
        "備考": "変更成功後、ボタンは disabled のまま維持する（二重送信防止）。"
    },
    {
        "id": "SCR-004", "name": "管理画面", "file": "admin.html",
        "url": "/admin.html",
        "遷移元": "管理者（is_admin=true）のログイン成功時",
        "概要": "管理者専用のユーザー管理画面。パスワード変更・ユーザー追加・ロック管理・有効化/無効化・削除が可能。",
        "入力項目": [
            ("管理者PW（新）", "password", "制限なし（1文字以上）", "必須のみ", "変更ボタン押下時"),
            ("追加ユーザーID", "text", "1〜20文字 / 半角英数字", "必須 / 文字種 / 文字数", "追加ボタン押下時"),
            ("追加パスワード", "password", "8〜32文字 / 半角英数記号", "必須 / 文字種 / 文字数", "追加ボタン押下時"),
        ],
        "ボタン": [
            ("管理者PW変更", "PUT /api/admin/password（長さ制限なし）"),
            ("ユーザー追加", "POST /api/users"),
            ("PW変更（モーダル）", "PUT /api/users/:id（password）"),
            ("PW変更要求", "PUT /api/users/:id（forcePasswordChange）"),
            ("ロックする", "PUT /api/users/:id（lock）"),
            ("ロック解除", "PUT /api/users/:id（resetLock）"),
            ("無効化 / 有効化", "PUT /api/users/:id（isActive）"),
            ("削除", "DELETE /api/users/:id"),
        ],
        "処理フロー": [
            "① GET /api/users → ユーザー一覧を取得して表示",
            "② 管理者（is_admin=true）は「管理者」バッジを表示",
            "③ 管理者行には削除・ロック・無効化ボタンを非表示",
            "④ 管理者パスワード変更カードは専用 PUT /api/admin/password を使用",
        ],
        "エラー表示": [
            ("追加エラー", "#add-alert-error", "赤背景", "API エラー時"),
            ("一覧操作エラー", "#list-alert-error", "赤背景", "API エラー時"),
            ("管理者PWエラー", "#admin-pw-alert-error", "赤背景", "API エラー時"),
        ],
        "備考": "ユーザー一覧は管理者バッジ（badge-admin / 青）/ 有効（badge-active / 緑）/ ロック中（badge-locked / 赤）/ 無効（badge-inactive / グレー）で色分け。"
    },
]

USER_STORIES = [
    ("US-01","一般ユーザー","高",
     "社員として、正しい ID とパスワードでログインしてホーム画面を表示したい",
     "Given user001 が有効状態で PW 期限内で存在する\n"
     "When ログイン画面で user001 / P@ssword を入力してログインボタンを押す\n"
     "Then ホーム画面が表示され「user001 さん」とログイン日時が表示される",
     "tests/e2e/login.spec.js", "SC-01"),
    ("US-02","管理者","高",
     "管理者として、admin / root でログインすると管理画面に遷移したい",
     "Given admin ユーザーが is_admin=true で存在する\n"
     "When ログイン画面で admin / root を入力してログインボタンを押す\n"
     "Then /admin.html へリダイレクトされ管理画面が表示される",
     "tests/e2e/login.spec.js", "SC-ADM-01（新規）"),
    ("US-03","一般ユーザー","高",
     "社員として、誤ったパスワードを入力したとき残り試行回数を確認したい",
     "Given user001 が有効状態で存在する\n"
     "When 誤ったパスワードでログインする（1回目）\n"
     "Then 「あと4回失敗するとロックされます」が表示される",
     "tests/e2e/login.spec.js", "SC-06"),
    ("US-04","一般ユーザー","最高",
     "社員として、5回連続でパスワードを間違えたとき、アカウントがロックされることを認識したい",
     "Given e2elock が有効状態で存在する\n"
     "When 誤ったパスワードでログインを5回繰り返す\n"
     "Then 5回目にロックメッセージが表示され、正しいパスワードでもログインできなくなる",
     "tests/e2e/login.spec.js", "SC-07"),
    ("US-05","一般ユーザー","高",
     "社員として、パスワードが期限切れのとき変更画面に案内されたい",
     "Given e2eexpired の PW 変更日が 91 日以上前\n"
     "When e2eexpired / 正しい PW でログインする\n"
     "Then パスワード変更画面が表示される",
     "tests/e2e/login.spec.js", "SC-09"),
    ("US-06","一般ユーザー","高",
     "社員として、パスワード変更画面で新しいパスワードを設定してログイン画面に戻りたい",
     "Given PW 変更画面が表示されている\n"
     "When 新 PW「NewPass1!」と確認 PW「NewPass1!」を入力して変更ボタンを押す\n"
     "Then 成功メッセージが表示され、2秒後にログイン画面へ遷移する",
     "tests/e2e/login.spec.js", "SC-10"),
    ("US-07","一般ユーザー","中",
     "社員として、ログアウトするとログイン画面に戻り入力欄がクリアされることを確認したい",
     "Given user001 でログイン済み\n"
     "When ログアウトボタンを押す\n"
     "Then ログイン画面が表示され、ユーザーID・PW欄が空になる",
     "tests/e2e/login.spec.js", "SC-11"),
    ("US-08","一般ユーザー","高",
     "社員として、ユーザーID が未入力のまま送信したときエラーを知りたい",
     "Given ログイン画面が表示されている\n"
     "When ユーザーIDを空のままログインボタンを押す\n"
     "Then 「ユーザーIDを入力してください」が #err-userid に表示される",
     "tests/e2e/login.spec.js", "SC-02"),
    ("US-09","一般ユーザー","高",
     "社員として、パスワードが 7 文字（最小値-1）のとき文字数エラーを知りたい",
     "Given ログイン画面が表示されている\n"
     "When PW欄に 7文字のパスワードを入力しフォーカスを外す\n"
     "Then 「パスワードは8〜32文字で入力してください」が表示される",
     "tests/e2e/login.spec.js", "SC-05"),
    ("US-10","管理者","高",
     "管理者として、新しい社員アカウントを安全に作成したい",
     "Given 管理画面が表示されている\n"
     "When 未使用ID と有効なPW を入力して追加ボタンを押す\n"
     "Then 成功メッセージが表示され、ユーザー一覧に新ユーザーが追加される",
     "tests/e2e/admin.spec.js", "SC-A01"),
    ("US-11","管理者","高",
     "管理者として、ロックされた社員のアカウントをすぐに解除したい",
     "Given e2elocktest がロック状態で存在する\n"
     "When ロック解除ボタンをクリックする\n"
     "Then badge-active バッジが表示され、失敗カウントが 0 にリセットされる",
     "tests/e2e/admin.spec.js", "SC-A06"),
    ("US-12","管理者","高",
     "管理者として、管理者パスワードを専用フォームから変更したい",
     "Given 管理画面が表示されている\n"
     "When 管理者PW変更カードに「newadminpw」を入力して変更ボタンを押す\n"
     "Then 成功メッセージが表示され、newadminpw で admin ログインが成功する",
     "tests/e2e/admin.spec.js", "SC-ADM-02（新規）"),
    ("US-13","管理者","高",
     "管理者として、退職者のアカウントを削除または無効化したい",
     "Given e2edelete / e2elocktest が有効状態で存在する\n"
     "When 削除ボタン または 無効化ボタンをクリックする\n"
     "Then 一覧から消える（削除）/ badge-inactive バッジが表示される（無効化）",
     "tests/e2e/admin.spec.js", "SC-A09 / SC-A11"),
]

TEST_CASES = [
    # (ID, 名前, 区分, 前提条件, 検証手順, 期待値, 自動化手段, 対応コード, 優先度)
    ("TC-E01","一般ユーザー正常ログイン","正常系",
     "・user001 が有効状態で存在する\n・PW 有効期限が 90 日以内\n・ブラウザでログイン画面を開いている",
     "1. ユーザーID欄に「user001」を入力\n2. パスワード欄に「P@ssword」を入力\n3. ログインボタンをクリック",
     "1. ホーム画面に切り替わる\n2. ホーム画面のユーザー名表示が「user001 さん」になる\n3. ログイン成功メッセージが表示される",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-01", "最高"),

    ("TC-E02","管理者ログイン → 管理画面遷移","正常系",
     "・admin ユーザーが is_admin=true で存在する\n・ブラウザでログイン画面を開いている",
     "1. ユーザーID欄に「admin」を入力\n2. パスワード欄に「root」を入力\n3. ログインボタンをクリック",
     "1. URL が /admin.html になる\n2. ページタイトルに「ユーザー管理」が含まれる\n3. ユーザー一覧が表示される",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-ADM-01", "最高"),

    ("TC-E03","ユーザーID未入力バリデーション","異常系",
     "・ブラウザでログイン画面を開いている",
     "1. ユーザーID欄を空のままにする\n2. パスワード欄に「P@ssword」を入力\n3. ログインボタンをクリック",
     "1. ユーザーIDエラーメッセージが表示される\n2. エラー文言が「ユーザーIDを入力してください」\n3. ユーザーID入力欄がエラー表示になる\n4. API 呼び出しは発生しない",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-02", "高"),

    ("TC-E04","パスワード未入力バリデーション","異常系",
     "・ブラウザでログイン画面を開いている",
     "1. ユーザーID欄に「user001」を入力\n2. パスワード欄を空のままにする\n3. ログインボタンをクリック",
     "1. パスワードエラーメッセージが表示される\n2. エラー文言が「パスワードを入力してください」\n3. API 呼び出しは発生しない",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-03", "高"),

    ("TC-E05","記号含むユーザーID バリデーション","異常系",
     "・ブラウザでログイン画面を開いている",
     "1. ユーザーID欄に「user_invalid」を入力\n2. フォーカスをパスワード欄に移す（blur）",
     "1. ユーザーIDエラーメッセージが表示される\n2. エラー文言が「ユーザーIDは半角英数字で入力してください」",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-04", "高"),

    ("TC-E06","パスワード 7文字（最小値-1）バリデーション","境界値",
     "・ブラウザでログイン画面を開いている",
     "1. パスワード欄に「P@ss001」（7文字）を入力\n2. フォーカスを外す（blur）",
     "1. パスワードエラーメッセージが表示される\n2. エラー文言が「パスワードは8〜32文字で入力してください」",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-05[BV]", "高"),

    ("TC-E07","パスワード 8文字（最小値）バリデーション","境界値",
     "・ブラウザでログイン画面を開いている",
     "1. パスワード欄に「P@ssw001」（8文字）を入力\n2. フォーカスを外す（blur）",
     "1. パスワードエラーメッセージが表示されない\n2. エラーメッセージが表示されない",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-05[BV-min]", "高"),

    ("TC-E08","認証失敗 - 残り試行回数表示（1回目）","異常系",
     "・user001 が有効状態で存在する（ログイン失敗回数がゼロの状態）",
     "1. ユーザーID「user001」/ パスワード「wrongpass」でログイン",
     "1. 認証エラーメッセージが表示される\n2. テキストに「あと4回失敗するとロックされます」が含まれる\n3. ロックエラー表示にはなっていない",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-06", "高"),

    ("TC-E09","5回連続失敗でアカウントロック","境界値",
     "・e2elock が有効状態で存在する（ログイン失敗回数がゼロの状態）",
     "1. e2elock / 誤PW でログイン（1回目）→ 残り4回確認\n"
     "2. e2elock / 誤PW でログイン（2回目）→ 残り3回確認\n"
     "3. e2elock / 誤PW でログイン（3回目）→ 残り2回確認\n"
     "4. e2elock / 誤PW でログイン（4回目）→ 残り1回確認\n"
     "5. e2elock / 誤PW でログイン（5回目）",
     "1. 1〜4回目：認証エラーメッセージに残り回数（4,3,2,1）が表示される\n"
     "2. 5回目：認証エラーメッセージがロックエラー表示になる\n"
     "3. 5回目：ロックメッセージが表示される",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-07", "最高"),

    ("TC-E10","ロック中のログイン拒否","異常系",
     "・e2elock がロック状態で存在する（locked_until が未来の日時）",
     "1. e2elock / 正しいパスワードでログインを試みる",
     "1. 認証エラーメッセージがロックエラー表示になる\n2. ロック中メッセージが表示される\n3. ホーム画面へは遷移しない",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-08", "最高"),

    ("TC-E11","PW期限切れ → 変更画面遷移","異常系",
     "・e2eexpired が有効状態で存在する（password_changed_at が 91 日以上前）",
     "1. e2eexpired / 正しいパスワードでログインする",
     "1. パスワード変更画面に切り替わる\n2. ログイン画面は非表示になる",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-09", "高"),

    ("TC-E12","PW変更成功と自動遷移","正常系",
     "・パスワード変更画面が表示されている状態",
     "1. 新しいパスワード欄に「NewPass1!」を入力\n"
     "2. 確認パスワード欄に「NewPass1!」を入力\n"
     "3. 「パスワードを変更する」ボタンをクリック",
     "1. パスワード変更成功メッセージが表示される\n"
     "2. 「パスワードを変更する」ボタンが非活性になる\n"
     "3. 2秒後にログイン画面が表示される",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-10", "高"),

    ("TC-E13","PW変更 - 確認PW不一致エラー","異常系",
     "・PW変更画面が表示されている状態",
     "1. 新しいパスワード欄に「NewPass1!」を入力\n"
     "2. 確認パスワード欄に「DifferentPass1!」を入力\n"
     "3. 「パスワードを変更する」ボタンをクリック",
     "1. 確認パスワードエラーメッセージに「パスワードが一致しません」が表示される\n"
     "2. パスワード変更APIは呼び出されない",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-12", "高"),

    ("TC-E14","ログアウトとフォームクリア","正常系",
     "・user001 でログイン済み（ホーム画面表示中）",
     "1. ログアウトボタンをクリックする",
     "1. ログイン画面が表示される\n"
     "2. ユーザーID入力欄が空になる\n"
     "3. パスワード入力欄が空になる\n"
     "4. 認証エラーメッセージは非表示のまま",
     "Playwright E2E", "tests/e2e/login.spec.js > SC-11", "高"),

    ("TC-A01","管理者：ユーザー追加成功","正常系",
     "・管理画面が表示されている\n・e2enew001 が DB に存在しない",
     "1. ユーザーID欄に「e2enew001」を入力\n"
     "2. パスワード欄に「NewPass1!」を入力\n"
     "3. 「追加」ボタンをクリック",
     "1. ユーザー追加成功アラートが表示される\n"
     "2. ユーザー一覧に e2enew001 が追加される\n"
     "3. e2enew001 に「有効」バッジが表示される",
     "Playwright E2E", "tests/e2e/admin.spec.js > SC-A01", "高"),

    ("TC-A02","管理者：ID重複エラー","異常系",
     "・user001 が DB に存在する",
     "1. ユーザーID欄に「user001」を入力\n"
     "2. パスワード欄に「NewPass1!」を入力\n"
     "3. 「追加」ボタンをクリック",
     "1. ユーザー追加エラーアラートが表示される\n"
     "2. エラーメッセージに「すでに使用されています」が含まれる",
     "Playwright E2E", "tests/e2e/admin.spec.js > SC-A02", "高"),

    ("TC-A03","管理者：手動ロック","正常系",
     "・e2elocktest が有効状態で存在する",
     "1. e2elocktest の「ロックする」ボタンをクリック\n"
     "2. 確認ダイアログで「OK」をクリック",
     "1. 成功メッセージが表示される\n"
     "2. e2elocktest のバッジが「ロック中」になる",
     "Playwright E2E", "tests/e2e/admin.spec.js > SC-A05", "高"),

    ("TC-A04","管理者：ロック解除","正常系",
     "・e2elocktest がロック状態で存在する",
     "1. e2elocktest の「ロック解除」ボタンをクリック",
     "1. 成功メッセージが表示される\n"
     "2. e2elocktest のバッジが「有効」になる",
     "Playwright E2E", "tests/e2e/admin.spec.js > SC-A06", "高"),

    ("TC-A05","管理者：管理者パスワード変更","正常系",
     "・管理画面が表示されている",
     "1. 管理者PW変更カードの入力欄に「newadminpw」を入力\n"
     "2. 「変更する」ボタンをクリック\n"
     "3. ログイン画面で admin / newadminpw でログインする",
     "1. 管理者パスワード変更成功アラートが表示される\n"
     "2. admin / newadminpw でログインに成功する\n"
     "3. 管理画面へリダイレクトされる",
     "Playwright E2E", "tests/e2e/admin.spec.js > SC-ADM-02", "最高"),

    ("TC-A06","管理者：管理者バッジ表示・保護確認","正常系",
     "・管理画面が表示されている",
     "1. ユーザー一覧を確認する",
     "1. admin 行に「管理者」バッジが表示される\n"
     "2. admin 行に削除・ロック・無効化ボタンが表示されない",
     "Playwright E2E", "tests/e2e/admin.spec.js > SC-ADM-03", "高"),

    ("TC-A07","管理者：ユーザー無効化","正常系",
     "・e2elocktest が有効状態で存在する",
     "1. e2elocktest の「無効化」ボタンをクリック\n"
     "2. 確認ダイアログで「OK」をクリック",
     "1. 成功メッセージが表示される\n"
     "2. e2elocktest のバッジが「無効」になる",
     "Playwright E2E", "tests/e2e/admin.spec.js > SC-A09", "中"),

    ("TC-P01","POST /api/auth/login - 正常ログイン","正常系（API）",
     "・user001 が有効状態で存在する（PW期限内）",
     "1. POST /api/auth/login\n   body: { userId: 'user001', password: 'P@ssword' }",
     "1. HTTP 200\n2. body.data.userId === 'user001'\n3. body.data.isAdmin === false",
     "Playwright API", "tests/auth.spec.js > TC-L01", "最高"),

    ("TC-P02","POST /api/auth/login - 管理者ログイン","正常系（API）",
     "・admin が is_admin=true で存在する",
     "1. POST /api/auth/login\n   body: { userId: 'admin', password: 'root' }",
     "1. HTTP 200\n2. body.data.isAdmin === true",
     "Playwright API", "tests/auth.spec.js > TC-ADM-L01", "最高"),

    ("TC-P03","POST /api/auth/login - PW不一致（1回目）","異常系（API）",
     "・user001 が有効状態（ログイン失敗回数がゼロの状態）",
     "1. POST /api/auth/login\n   body: { userId: 'user001', password: 'wrong' }",
     "1. HTTP 401\n2. body.error.code === 'AUTH_FAILED'\n3. body.error.remainingAttempts === 4",
     "Playwright API", "tests/auth.spec.js > TC-L03", "高"),

    ("TC-P04","POST /api/auth/login - 5回失敗でロック","境界値（API）",
     "・user001 が有効状態（ログイン失敗回数がゼロの状態）",
     "1〜4回: POST /api/auth/login（誤PW）→ 401 確認\n"
     "5回目: POST /api/auth/login（誤PW）",
     "5回目:\n1. HTTP 403\n2. body.error.code === 'ACCOUNT_LOCKED'\n3. body.error.lockedUntil が設定される",
     "Playwright API", "tests/auth.spec.js > TC-L04〜L08", "最高"),

    ("TC-P05","POST /api/auth/logout - セッション無効化","正常系（API）",
     "・user001 でログイン済み（有効なセッション）",
     "1. POST /api/auth/logout\n2. GET /api/auth/status",
     "1. ログアウト: HTTP 200\n2. status 確認: HTTP 401",
     "Playwright API", "tests/auth.spec.js > TC-LO01/LO02", "高"),

    ("TC-P06","PUT /api/admin/password - 管理者PW変更","正常系（API）",
     "・admin ユーザーが is_admin=true で存在する",
     "1. PUT /api/admin/password\n   body: { password: 'newpw' }\n"
     "2. POST /api/auth/login（admin / newpw）",
     "1. HTTP 200\n2. body.status === 'success'\n3. 新PWでのログインが成功（HTTP 200）",
     "Playwright API", "tests/auth.spec.js > TC-ADM-PW01", "最高"),

    ("TC-U01","isPasswordExpired - 89日前（境界値-1）","境界値（Unit）",
     "・utils.js の isPasswordExpired 関数",
     "1. isPasswordExpired(89日前の Date) を呼び出す",
     "1. 戻り値が false",
     "Jest Unit", "tests/unit/utils.test.js > BV-03", "中"),

    ("TC-U02","isPasswordExpired - 90日前（境界値）","境界値（Unit）",
     "・utils.js の isPasswordExpired 関数",
     "1. isPasswordExpired(90日前の Date) を呼び出す",
     "1. 戻り値が true",
     "Jest Unit", "tests/unit/utils.test.js > BV-04", "高"),

    ("TC-U03","isLocked - ロック中（1分後）","境界値（Unit）",
     "・utils.js の isLocked 関数",
     "1. isLocked(1分後の Date) を呼び出す",
     "1. 戻り値が true",
     "Jest Unit", "tests/unit/utils.test.js > BV-15", "高"),

    ("TC-U04","isLocked - ロック期限切れ（1分前）","境界値（Unit）",
     "・utils.js の isLocked 関数",
     "1. isLocked(1分前の Date) を呼び出す",
     "1. 戻り値が false",
     "Jest Unit", "tests/unit/utils.test.js > BV-13", "高"),

    ("TC-S01","SQLインジェクション耐性（ユーザーID）","セキュリティ",
     "・DB が稼働している",
     "1. POST /api/auth/login\n   body: { userId: \"' OR '1'='1\", password: 'any' }",
     "1. HTTP 400 または 401\n2. 全ユーザーが返却されない\n3. DB が改ざんされない（GET /api/users で確認）",
     "Playwright API", "tests/auth.spec.js > TC-SEC-01", "最高"),

    ("TC-S02","ログアウト後の保護API拒否","セキュリティ",
     "・ログアウト済み（セッションなし）",
     "1. GET /api/auth/status を送信する",
     "1. HTTP 401\n2. body.error.code === 'UNAUTHORIZED'",
     "Playwright API", "tests/auth.spec.js > TC-ST02", "高"),
]

ACCOUNT_STATE_TRANSITIONS = [
    ("有効（active）","連続5回認証失敗","ロック中（locked）","ログイン連続失敗回数が5回に達し、ロック期限が設定される"),
    ("有効（active）","管理者が手動ロック","ロック中（locked）","ロック期限が永続（9999-12-31）に設定される"),
    ("有効（active）","管理者が無効化","無効（inactive）","アカウントが無効化される"),
    ("有効（active）","PW未変更から90日経過","PW期限切れ（expired_pw）","次回ログイン時にチェック、状態は有効のまま"),
    ("ロック中（locked）","30分後（自動）","有効（active）","ロック期限が過去になり、失敗回数がリセットされる"),
    ("ロック中（locked）","管理者がロック解除","有効（active）","ロック期限が解除され、失敗回数がリセットされる"),
    ("無効（inactive）","管理者が有効化","有効（active）","アカウントが有効化される"),
    ("PW期限切れ","ユーザーがPW変更","有効（active）","パスワード変更日時が現在時刻に更新される"),
    ("PW期限切れ","管理者がPW変更要求","有効（active）","次回ログイン時にパスワード変更画面へ誘導される"),
    ("管理者（admin）","管理者PW変更フォームで変更","管理者（admin）","管理者フラグは変化しない。パスワードのみ更新される"),
]

SCREEN_TRANSITIONS = [
    ("ログイン画面","一般ユーザーログイン成功","ホーム画面","HTTP 200 / isAdmin=false"),
    ("ログイン画面","管理者ログイン成功","管理画面（/admin.html）","HTTP 200 / isAdmin=true → window.location.href"),
    ("ログイン画面","PW期限切れ（HTTP 422）","PW変更画面","PASSWORD_EXPIRED コード"),
    ("ログイン画面","認証失敗（HTTP 401）","ログイン画面（エラー表示）","残り試行回数を表示"),
    ("ログイン画面","アカウントロック（HTTP 403）","ログイン画面（ロックエラー表示）","ロック残り時間を表示"),
    ("ホーム画面","ログアウトボタン押下","ログイン画面","フォームクリア・セッション破棄"),
    ("PW変更画面","PW変更成功（HTTP 200）","ログイン画面（2秒後）","成功メッセージ → ボタン disabled → 2秒タイマー"),
    ("PW変更画面","バリデーションエラー","PW変更画面（エラー表示）","API呼び出しなし"),
    ("管理画面","各ユーザー管理操作","管理画面（一覧更新）","PUT/DELETE 後に loadUsers() 再実行"),
]

# ──────────────────────────────────────────────
# Excel 生成
# ──────────────────────────────────────────────

def generate_excel():
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # ── シート1: 画面仕様サマリー ──
    ws = wb.create_sheet("画面仕様書")
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A3"
    xl_title(ws, "画面仕様書 — 社内Webシステム ログイン・ユーザー管理機能", 8)
    hdrs = ["画面ID","画面名","ファイル","URL / 遷移元","入力項目","ボタン","処理フロー概要","エラー表示","備考"]
    for i,h in enumerate(hdrs,1):
        xl_cell(ws,2,i,h,bold=True,bg="2E75B6",h="center",v="center")
        ws.cell(row=2,column=i).font = Font(name="メイリオ",bold=True,color="FFFFFF",size=9)
        ws.cell(row=2,column=i).border = thin_border()
    ws.row_dimensions[2].height=20
    stripe=["FFFFFF","EEF3FA"]
    for ri,sc in enumerate(SCREEN_SPECS):
        r=ri+3; bg=stripe[ri%2]
        inputs="\n".join(f"・{x[0]}（{x[2]}）" for x in sc["入力項目"]) or "—"
        buttons="\n".join(f"・{b[0]}" for b in sc["ボタン"])
        flow="\n".join(sc["処理フロー"])
        errs="\n".join(f"・{e[0]}" for e in sc["エラー表示"]) or "—"
        vals=[sc["id"],sc["name"],sc["file"],f"{sc['url']}\n遷移元: {sc['遷移元']}",inputs,buttons,flow,errs,sc["備考"]]
        for ci,v in enumerate(vals,1):
            xl_cell(ws,r,ci,v,bg=bg)
        ws.row_dimensions[r].height=80
    widths=[2.5,4.0,3.5,5.5,6.0,5.5,10.0,5.5,5.0]
    for i,w in enumerate(widths,1): set_w(ws,i,w)

    # ── シート2: 画面遷移表 ──
    ws2=wb.create_sheet("画面遷移")
    ws2.sheet_view.showGridLines=False
    ws2.freeze_panes="A3"
    xl_title(ws2,"画面遷移表 — 社内Webシステム ログイン・ユーザー管理機能",4)
    for i,h in enumerate(["遷移元画面","トリガー","遷移先画面","条件 / 備考"],1):
        xl_cell(ws2,2,i,h,bold=True,bg="2E75B6",h="center",v="center")
        ws2.cell(row=2,column=i).font=Font(name="メイリオ",bold=True,color="FFFFFF",size=9)
        ws2.cell(row=2,column=i).border=thin_border()
    ws2.row_dimensions[2].height=20
    src_colors={"ログイン画面":"DEEBF7","ホーム画面":"E2EFDA","PW変更画面":"FFF2CC","管理画面（/admin.html）":"EDE7F6"}
    for ri,(src,trig,dst,cond) in enumerate(SCREEN_TRANSITIONS):
        r=ri+3; bg=src_colors.get(src,"FFFFFF")
        for ci,v in enumerate([src,trig,dst,cond],1):
            xl_cell(ws2,r,ci,v,bg=bg if ci<=2 else "FFFFFF")
        ws2.row_dimensions[r].height=28
    for i,w in enumerate([5.0,7.0,7.0,10.0],1): set_w(ws2,i,w)

    # ── シート3: アカウント状態遷移表 ──
    ws3=wb.create_sheet("アカウント状態遷移")
    ws3.sheet_view.showGridLines=False
    ws3.freeze_panes="A3"
    xl_title(ws3,"アカウント状態遷移表 — 社内Webシステム ログイン機能",4)
    for i,h in enumerate(["遷移前状態","トリガー","遷移後状態","DB変化 / 備考"],1):
        xl_cell(ws3,2,i,h,bold=True,bg="2E75B6",h="center",v="center")
        ws3.cell(row=2,column=i).font=Font(name="メイリオ",bold=True,color="FFFFFF",size=9)
        ws3.cell(row=2,column=i).border=thin_border()
    ws3.row_dimensions[2].height=20
    st_colors={"有効（active）":"E2EFDA","ロック中（locked）":"FED7D7","無効（inactive）":"E2E8F0","PW期限切れ（expired_pw）":"FEEBC8","管理者（admin）":"BEE3F8"}
    for ri,(before,trig,after,note) in enumerate(ACCOUNT_STATE_TRANSITIONS):
        r=ri+3; bg=st_colors.get(before,"FFFFFF"); abg=st_colors.get(after,"FFFFFF")
        xl_cell(ws3,r,1,before,bg=bg,bold=True)
        xl_cell(ws3,r,2,trig)
        xl_cell(ws3,r,3,after,bg=abg,bold=True)
        xl_cell(ws3,r,4,note)
        ws3.row_dimensions[r].height=28
    for i,w in enumerate([5.5,7.0,5.5,11.0],1): set_w(ws3,i,w)

    # ── シート4: ユーザーストーリー ──
    ws4=wb.create_sheet("ユーザーストーリー")
    ws4.sheet_view.showGridLines=False
    ws4.freeze_panes="A3"
    xl_title(ws4,"ユーザーストーリー — 自動化シナリオ前提（Given / When / Then）",7)
    for i,h in enumerate(["US-ID","ロール","優先度","ストーリー","Given / When / Then（自動化シナリオ）","対応テストファイル","対応テストID"],1):
        xl_cell(ws4,2,i,h,bold=True,bg="2E75B6",h="center",v="center")
        ws4.cell(row=2,column=i).font=Font(name="メイリオ",bold=True,color="FFFFFF",size=9)
        ws4.cell(row=2,column=i).border=thin_border()
    ws4.row_dimensions[2].height=20
    prio_c={"最高":"C00000","高":"C55A11","中":"BF8F00"}
    stripe=["FFFFFF","EEF3FA"]
    for ri,us in enumerate(USER_STORIES):
        r=ri+3; bg=stripe[ri%2]
        usid,role,prio,story,gwt,f,tid=us
        xl_cell(ws4,r,1,usid,bold=True,bg=bg,h="center")
        xl_cell(ws4,r,2,role,bg=bg,h="center")
        pc=ws4.cell(row=r,column=3,value=prio)
        pc.font=Font(name="メイリオ",bold=True,color="FFFFFF",size=9)
        pc.fill=hex_fill(prio_c.get(prio,"595959"))
        pc.alignment=wrap("center","center"); pc.border=thin_border()
        xl_cell(ws4,r,4,story,bg=bg)
        xl_cell(ws4,r,5,gwt)
        xl_cell(ws4,r,6,f)
        xl_cell(ws4,r,7,tid)
        ws4.row_dimensions[r].height=65
    for i,w in enumerate([2.0,3.0,2.0,9.0,14.0,7.0,6.0],1): set_w(ws4,i,w)

    # ── シート5: テストケース ──
    ws5=wb.create_sheet("テストケース")
    ws5.sheet_view.showGridLines=False
    ws5.freeze_panes="A3"
    xl_title(ws5,"テストケース — 前提条件 / 検証手順 / 期待値を分離（自動化前提）",9)
    tc_hdrs=["テストケースID","テスト名","テスト区分","前提条件（Preconditions）",
             "検証手順（Steps）","期待値（Expected Results）","自動化手段","対応コード","優先度"]
    for i,h in enumerate(tc_hdrs,1):
        xl_cell(ws5,2,i,h,bold=True,bg="1F4E79",h="center",v="center")
        ws5.cell(row=2,column=i).font=Font(name="メイリオ",bold=True,color="FFFFFF",size=9)
        ws5.cell(row=2,column=i).border=thin_border()
    ws5.row_dimensions[2].height=22
    type_bg={"正常系":"E2EFDA","異常系":"FFF2CC","境界値":"DEEBF7","境界値（API）":"DEEBF7",
             "境界値（Unit）":"EDE7F6","セキュリティ":"FED7D7","正常系（API）":"E2EFDA",
             "異常系（API）":"FFF2CC"}
    prio_c2={"最高":"C00000","高":"C55A11","中":"BF8F00","低":"375623"}
    for ri,tc in enumerate(TEST_CASES):
        r=ri+3
        tcid,name,ttype,pre,steps,exp,tool,code,prio=tc
        tbg=type_bg.get(ttype,"FFFFFF")
        xl_cell(ws5,r,1,tcid,bold=True,bg=tbg,h="center")
        xl_cell(ws5,r,2,name,bold=True,bg=tbg)
        xl_cell(ws5,r,3,ttype,bg=tbg,h="center")
        xl_cell(ws5,r,4,pre)
        xl_cell(ws5,r,5,steps)
        xl_cell(ws5,r,6,exp)
        xl_cell(ws5,r,7,tool,h="center")
        xl_cell(ws5,r,8,code)
        pc=ws5.cell(row=r,column=9,value=prio)
        pc.font=Font(name="メイリオ",bold=True,color="FFFFFF",size=9)
        pc.fill=hex_fill(prio_c2.get(prio,"595959"))
        pc.alignment=wrap("center","center"); pc.border=thin_border()
        ws5.row_dimensions[r].height=80
    for i,w in enumerate([2.8,5.5,3.0,8.0,10.0,9.0,3.5,7.0,2.0],1): set_w(ws5,i,w)
    ws5.auto_filter.ref=f"A2:I{len(TEST_CASES)+2}"

    path=OUT_DIR/"specs-and-testcases.xlsx"
    wb.save(str(path)); print(f"Excel 生成完了 → {path}")

# ──────────────────────────────────────────────
# Word 生成
# ──────────────────────────────────────────────

def generate_word():
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # タイトル
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("テスト設計ドキュメント"); r.bold=True; r.font.size=Pt(20)
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run("社内Webシステム ログイン・ユーザー管理機能（管理者機能追加版）")
    r2.font.size=Pt(12); r2.font.color.rgb=RGBColor(0x2E,0x75,0xB6)
    doc.add_paragraph()
    w_table(doc,["項目","内容"],
        [("対象システム","社内Webシステム ログイン・ユーザー管理機能（login-app）"),
         ("作成日","2026-05-14"),("バージョン","v2.0（管理者機能追加）"),("作成者","UedaKazuki")],
        [4.0,13.0],hdr_bg="2E75B6")
    doc.add_paragraph()

    # ══════════════════════════════
    # 1. 画面仕様書
    # ══════════════════════════════
    w_heading(doc,"1. 画面仕様書",1)

    for sc in SCREEN_SPECS:
        w_heading(doc,f"{sc['id']} {sc['name']}（{sc['file']}）",2)
        w_table(doc,["項目","内容"],
            [("URL / 遷移元",f"{sc['url']} ／ {sc['遷移元']}"),("概要",sc["概要"])],
            [3.5,13.5],hdr_bg="595959")
        doc.add_paragraph()

        if sc["入力項目"]:
            w_heading(doc,"入力項目",3)
            w_table(doc,["フィールド名","入力タイプ","制約","バリデーション","タイミング"],
                sc["入力項目"],[3.5,3.5,4.0,4.0,3.0],hdr_bg="2E75B6")
            doc.add_paragraph()

        w_heading(doc,"ボタン",3)
        w_table(doc,["ボタン名","動作・API"],sc["ボタン"],[3.5,13.5],hdr_bg="2E75B6")
        doc.add_paragraph()

        w_heading(doc,"処理フロー",3)
        for line in sc["処理フロー"]:
            p=doc.add_paragraph(line,style="List Bullet")
            p.runs[0].font.size=Pt(9)

        if sc["エラー表示"]:
            doc.add_paragraph()
            w_heading(doc,"エラー表示",3)
            w_table(doc,["種別","表示位置","スタイル","タイミング"],
                sc["エラー表示"],[3.5,4.0,4.0,5.5],hdr_bg="2E75B6")

        if sc["備考"]:
            doc.add_paragraph()
            w_para(doc,f"備考: {sc['備考']}",size=9,italic=True)
        doc.add_paragraph()

    # ══════════════════════════════
    # 2. 状態遷移図
    # ══════════════════════════════
    w_heading(doc,"2. 状態遷移図",1)

    w_heading(doc,"2-1. 画面遷移",2)
    w_para(doc,
        "各画面間の遷移トリガーと条件を示す。SPA のため SCR-001〜003 は同一ページ内の表示切替。"
        "SCR-004（管理画面）のみ別 HTML ファイルへのリダイレクト。",size=9)
    doc.add_paragraph()

    # 画面遷移テキスト図
    p=doc.add_paragraph()
    run=p.add_run(
        "【ログイン画面 SCR-001】\n"
        "  ├─ 一般ユーザーログイン成功（HTTP 200 / isAdmin=false）─────────────► 【ホーム画面 SCR-002】\n"
        "  ├─ 管理者ログイン成功（HTTP 200 / isAdmin=true）──────────────────────► 【管理画面 SCR-004 /admin.html】\n"
        "  ├─ PW期限切れ（HTTP 422 PASSWORD_EXPIRED）────────────────────────────► 【PW変更画面 SCR-003】\n"
        "  ├─ 認証失敗（HTTP 401 AUTH_FAILED）────────────────────────────────────► 【SCR-001 エラー表示】\n"
        "  └─ ロック中（HTTP 403 ACCOUNT_LOCKED）─────────────────────────────────► 【SCR-001 ロックエラー表示】\n\n"
        "【ホーム画面 SCR-002】\n"
        "  └─ ログアウトボタン押下 ────────────────────────────────────────────────► 【ログイン画面 SCR-001】\n\n"
        "【PW変更画面 SCR-003】\n"
        "  ├─ PW変更成功（HTTP 200）→ 成功メッセージ → 2秒後 ─────────────────────► 【ログイン画面 SCR-001】\n"
        "  └─ バリデーションエラー ─────────────────────────────────────────────────► 【SCR-003 エラー表示】\n\n"
        "【管理画面 SCR-004】\n"
        "  └─ 各操作（追加/ロック/削除/PW変更） → GET /api/users で一覧更新 ─────► 【SCR-004 同一画面更新】"
    )
    run.font.name="Courier New"; run.font.size=Pt(8)
    doc.add_paragraph()

    w_heading(doc,"2-2. アカウント状態遷移",2)
    w_para(doc,"ユーザーアカウントが取りうる状態と、状態を変化させるトリガーを示す。",size=9)
    doc.add_paragraph()

    p2=doc.add_paragraph()
    run2=p2.add_run(
        "                       連続5回失敗 / 管理者手動ロック\n"
        "   ┌──────────────────────────────────────────────────────────────┐\n"
        "   │                                                              ▼\n"
        "【有効 active】 ─── 管理者が無効化 ─────────────────────────► 【無効 inactive】\n"
        "   │  ▲                                                              │\n"
        "   │  └─────────── 管理者が有効化 ────────────────────────────────────┘\n"
        "   │\n"
        "   ├─── PW未変更90日超 ──────► 【PW期限切れ expired_pw】\n"
        "   │                               │（有効状態は変わらず、次回ログイン時に判定）\n"
        "   │                               └── PW変更 ──────────────────────► 【有効 active】\n"
        "   │\n"
        "   └─────────────────────────► 【ロック中 locked】\n"
        "                                   │\n"
        "                                   ├── 30分後（自動）─────────────► 【有効 active】\n"
        "                                   └── 管理者が解除 ───────────────► 【有効 active】\n\n"
        "【管理者 admin（is_admin=true）】\n"
        "   ├── ロック / 無効化 / 削除 操作なし（管理画面から保護）\n"
        "   └── 管理者PWフォームからパスワードのみ変更可（長さ制限なし）"
    )
    run2.font.name="Courier New"; run2.font.size=Pt(8)
    doc.add_paragraph()

    w_heading(doc,"2-3. 状態遷移表",2)
    w_table(doc,["遷移前状態","トリガー","遷移後状態","DB変化 / 備考"],
        ACCOUNT_STATE_TRANSITIONS,
        [4.0,5.5,4.5,8.5],hdr_bg="2E75B6")
    doc.add_paragraph()

    # ══════════════════════════════
    # 3. ユーザーストーリー
    # ══════════════════════════════
    w_heading(doc,"3. ユーザーストーリー（自動化前提）",1)
    w_para(doc,
        "各ストーリーは Playwright（E2E / API）または Jest（Unit）で自動化することを前提に Given / When / Then 形式で記述。",
        size=9,italic=True)
    doc.add_paragraph()

    for us in USER_STORIES:
        usid,role,prio,story,gwt,f,tid=us
        w_heading(doc,f"{usid} : {story}",2)
        w_table(doc,["属性","内容"],
            [("ロール",role),("優先度",prio),("対応テスト",f"{f} / {tid}")],
            [3.0,14.0],hdr_bg="595959")
        doc.add_paragraph()
        p=doc.add_paragraph()
        run=p.add_run(gwt); run.font.name="Courier New"; run.font.size=Pt(9)
        doc.add_paragraph()

    # ══════════════════════════════
    # 4. テストケース
    # ══════════════════════════════
    w_heading(doc,"4. テストケース（前提条件 / 検証手順 / 期待値を分離）",1)
    w_para(doc,
        "各テストケースは 前提条件（Preconditions）/ 検証手順（Steps）/ 期待値（Expected Results）を明確に分離。"
        "詳細は Excel シート「テストケース」を参照。",size=9,italic=True)
    doc.add_paragraph()

    for tc in TEST_CASES:
        tcid,name,ttype,pre,steps,exp,tool,code,prio=tc
        w_heading(doc,f"{tcid} : {name}",2)
        w_table(doc,["属性","内容"],
            [("テスト区分",ttype),("優先度",prio),("自動化手段",tool),("対応コード",code)],
            [3.5,13.5],hdr_bg="595959")
        doc.add_paragraph()

        for label,content,bg_hex in [
            ("■ 前提条件（Preconditions）",pre,"DEEBF7"),
            ("■ 検証手順（Steps）",steps,"E2EFDA"),
            ("■ 期待値（Expected Results）",exp,"FFF2CC"),
        ]:
            lp=doc.add_paragraph()
            lr=lp.add_run(label); lr.bold=True; lr.font.size=Pt(9)
            cp=doc.add_paragraph()
            cr=cp.add_run(content); cr.font.size=Pt(9)
            cr.font.name="Meiryo"
            cp.paragraph_format.left_indent=Cm(0.8)

        doc.add_paragraph()

    path=OUT_DIR/"specs-and-testcases.docx"
    doc.save(str(path)); print(f"Word 生成完了 → {path}")

if __name__=="__main__":
    generate_excel()
    generate_word()
