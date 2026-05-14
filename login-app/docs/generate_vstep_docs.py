"""
VSTeP テスト設計書 — Word (.docx) + Excel (.xlsx) 生成スクリプト
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter
import copy
from pathlib import Path

OUT_DIR = Path(__file__).parent

# ─────────────────────────────────────────
# 共通データ定義
# ─────────────────────────────────────────

TEST_CONDITIONS = [
    # (ID, フレーム, グループ, 技法, コンディション, 入力/前提状態, 期待結果, 優先度)
    ("TC-F1-01","F1","正常認証","同値分割（有効）",
     "有効ユーザー + 正しい ID/PW でログイン成功する",
     "有効ユーザー + 正しい ID + 正しい PW",
     "200 → ホーム画面遷移・ユーザー名表示","最高"),
    ("TC-F1-02","F1","認証失敗","同値分割（無効）",
     "存在しない ID ではログインできない",
     "未登録 ID + 任意 PW","401 AUTH_FAILED","高"),
    ("TC-F1-03","F1","認証失敗","同値分割（無効）",
     "無効化されたユーザーはログインできない",
     "有効 ID + 正しい PW + is_active=false","401 AUTH_FAILED","高"),
    ("TC-F1-04","F1","認証失敗/残り回数","同値分割",
     "PW 不一致 1回目 → 残り4回を表示する",
     "有効 ID + 誤 PW（1回目）","401 + remainingAttempts: 4","高"),
    ("TC-F1-05","F1","ロック境界値","境界値（閾値-1）",
     "PW 不一致 4回目 → 残り1回を表示しロックしない",
     "有効 ID + 誤 PW（4回目）","401 + remainingAttempts: 1","高"),
    ("TC-F1-06","F1","ロック発動","境界値（閾値）",
     "PW 不一致 5回目 → ACCOUNT_LOCKED になる",
     "有効 ID + 誤 PW（5回目）","403 ACCOUNT_LOCKED","最高"),
    ("TC-F1-07","F1","ロック中認証","状態遷移（ロック中）",
     "ロック中は正しい PW でもログインできない",
     "ロック状態 + 正しい PW","403 ACCOUNT_LOCKED","最高"),
    ("TC-F1-08","F1","期限切れ","状態遷移",
     "PW 期限切れユーザーは変更画面へ遷移する",
     "password_changed_at = 90日以上前","422 PASSWORD_EXPIRED → PW 変更画面","高"),
    ("TC-F2-01","F2","バリデーション/ID","同値分割（無効）",
     "ユーザーID 未入力 → 必須エラーを表示する",
     "空文字（送信時）","「ユーザーIDを入力してください」","高"),
    ("TC-F2-02","F2","バリデーション/ID","同値分割（無効）",
     "記号を含む ID → 文字種エラーを表示する",
     "user_invalid（記号含む）","「半角英数字で入力してください」","高"),
    ("TC-F2-03","F2","バリデーション/ID","境界値（下限）",
     "ID 1文字（最小値）→ エラーなし",
     "a（1文字）","エラーなし","中"),
    ("TC-F2-04","F2","バリデーション/ID","境界値（上限）",
     "ID 20文字（最大値）→ エラーなし",
     "a × 20文字","エラーなし","中"),
    ("TC-F2-05","F2","バリデーション/ID","境界値（上限+1）",
     "ID 21文字（最大値+1）→ 長さエラーを表示する",
     "a × 21文字","「1〜20文字で入力してください」","中"),
    ("TC-F3-01","F3","バリデーション/PW","同値分割（無効）",
     "PW 未入力 → 必須エラーを表示する",
     "空文字","「パスワードを入力してください」","高"),
    ("TC-F3-02","F3","バリデーション/PW","境界値（下限-1）",
     "PW 7文字（最小値-1）→ 長さエラーを表示する",
     "P@ss001（7文字）","「8〜32文字で入力してください」","高"),
    ("TC-F3-03","F3","バリデーション/PW","境界値（下限）",
     "PW 8文字（最小値）→ エラーなし",
     "P@ssw001（8文字）","エラーなし","高"),
    ("TC-F3-04","F3","バリデーション/PW","境界値（上限）",
     "PW 32文字（最大値）→ エラーなし",
     "P@ssword + a×24（32文字）","エラーなし","中"),
    ("TC-F3-05","F3","バリデーション/PW","境界値（上限+1）",
     "PW 33文字（最大値+1）→ 長さエラーを表示する",
     "P@ssword + a×25（33文字）","「8〜32文字で入力してください」","中"),
    ("TC-F4-01","F4","ロック/残り回数","境界値（閾値-4）",
     "1回失敗時に残り4回を正しく表示する",
     "失敗累積1回","残り4回メッセージ","高"),
    ("TC-F4-02","F4","ロック/残り回数","境界値（閾値-1）",
     "4回失敗時に残り1回を表示しロックしない",
     "失敗累積4回","残り1回メッセージ / ロックなし","高"),
    ("TC-F4-03","F4","ロック発動","境界値（閾値）",
     "5回目失敗でロックが発動する",
     "失敗累積5回","ACCOUNT_LOCKED / auth-locked クラス付与","最高"),
    ("TC-F4-04","F4","ロック継続","状態遷移",
     "ロック中に再試行しても ACCOUNT_LOCKED になる",
     "ロック状態で任意 PW","ACCOUNT_LOCKED","最高"),
    ("TC-F5-01","F5","PW 期限","境界値（閾値-1）",
     "89日前の PW → 期限切れでない",
     "password_changed_at = 89日前","200 正常ログイン","中"),
    ("TC-F5-02","F5","PW 期限","境界値（閾値）",
     "90日前の PW → 期限切れと判定される",
     "password_changed_at = 90日前","422 PASSWORD_EXPIRED","高"),
    ("TC-F5-03","F5","PW 期限","境界値（閾値+1）",
     "91日前の PW → 期限切れと判定される",
     "password_changed_at = 91日前","422 PASSWORD_EXPIRED","中"),
    ("TC-F6-01","F6","PW変更/正常","同値分割（有効）",
     "新 PW = 確認 PW → 変更成功する",
     "一致する有効 PW × 2 / 期限切れセッション",
     "200 → 成功メッセージ → 2秒後ログイン画面","高"),
    ("TC-F6-02","F6","PW変更/不一致","同値分割（無効）",
     "新 PW ≠ 確認 PW → フィールドエラーを表示する",
     "異なる PW × 2","「パスワードが一致しません」","高"),
    ("TC-F6-03","F6","PW変更/未認証","状態遷移",
     "通常ログインセッションでは変更 API を使えない",
     "通常セッション + 有効 PW","401","高"),
    ("TC-F6-04","F6","PW変更/UI","UI動作確認",
     "変更成功後にボタンが disabled のまま維持される",
     "変更成功直後","#btn-change が disabled","中"),
    ("TC-F7-01","F7","セッション/認証済み","状態遷移",
     "ログイン済みセッションで status API が 200 を返す",
     "有効セッション","200 { userId }","高"),
    ("TC-F7-02","F7","セッション/未認証","状態遷移",
     "未認証では status API が 401 を返す",
     "セッションなし","401","高"),
    ("TC-F7-03","F7","ログアウト","状態遷移",
     "ログアウト後にセッションが無効化される",
     "ログイン済み → ログアウト → status API",
     "200（ログアウト成功）→ 401（セッション無効）","高"),
    ("TC-F7-04","F7","ログアウト/UI","UI動作確認",
     "ログアウト後に入力欄がクリアされる",
     "ログイン済みでログアウト","ログイン画面表示 / ID・PW 欄が空","高"),
    ("TC-F8-01","F8","ユーザー追加","同値分割（有効）",
     "新規ユーザーを正常に追加できる",
     "未使用 ID + 有効 PW","201 → 一覧に表示","高"),
    ("TC-F8-02","F8","ユーザー追加","同値分割（無効）",
     "同じ ID で重複追加するとエラーになる",
     "既存 ID + 任意 PW","400 DUPLICATE_USER","高"),
    ("TC-F8-03","F8","手動ロック","状態遷移（有効→ロック）",
     "管理者が有効ユーザーを手動ロックできる",
     "有効ユーザーをロック操作","badge-locked 表示","高"),
    ("TC-F8-04","F8","ロック解除","状態遷移（ロック→有効）",
     "管理者がロックを解除すると再ログイン可能になる",
     "ロックユーザーを解除操作",
     "badge-active 表示 / 失敗カウントリセット","高"),
    ("TC-F8-05","F8","無効化","状態遷移（有効→無効）",
     "管理者が有効ユーザーを無効化できる",
     "有効ユーザーを無効化操作","badge-inactive 表示","中"),
    ("TC-F8-06","F8","有効化","状態遷移（無効→有効）",
     "管理者が無効ユーザーを有効化できる",
     "無効ユーザーを有効化操作","badge-active 表示","中"),
    ("TC-F8-07","F8","ユーザー削除","同値分割（有効）",
     "存在するユーザーを削除すると一覧から消える",
     "存在するユーザーを削除操作","200 → 一覧から消える","中"),
    ("TC-SEC-01","—","セキュリティ","セキュリティ",
     "ユーザーID フィールドへの SQL インジェクションを拒否する",
     "' OR '1'='1（記号含む → バリデーションで拒否）",
     "400 または 401（DB 改ざんなし）","最高"),
    ("TC-SEC-02","—","セキュリティ","セキュリティ",
     "パスワードフィールドへの SQL インジェクションで認証突破できない",
     "' OR '1'='1; --",
     "401 AUTH_FAILED（DB 改ざんなし）","最高"),
    ("TC-SEC-03","—","セキュリティ","セキュリティ",
     "ログアウト後に保護 API へのアクセスを拒否する",
     "ログアウト後に GET /api/auth/status","401","高"),
]

FRAMES = [
    ("F1","ログイン認証処理",
     "POST /api/auth/login",
     "ユーザー状態（有効/無効/ロック/期限切れ）× ID・PW の正否",
     "HTTP ステータス・レスポンス内容・画面遷移先",
     "同値分割・状態遷移"),
    ("F2","ユーザーID バリデーション",
     "ログインフォーム #userid フィールド",
     "入力値の種類（空・文字種・文字数）× タイミング（blur/送信時）",
     "エラーメッセージの有無・内容・表示箇所",
     "同値分割・境界値分析"),
    ("F3","パスワード バリデーション",
     "ログインフォーム #password フィールド",
     "入力値の種類（空・文字種・文字数境界値）",
     "エラーメッセージの有無・内容",
     "同値分割・境界値分析"),
    ("F4","アカウントロック制御",
     "連続失敗カウント機能 / ロック発動ロジック",
     "累積失敗回数（1〜5回）",
     "残り試行回数の表示 / ロック発動の有無",
     "境界値分析・状態遷移"),
    ("F5","パスワード期限管理",
     "isPasswordExpired() 関数 / ログイン時の期限チェック",
     "password_changed_at からの経過日数",
     "正常ログイン成功 or PW 変更画面への遷移",
     "境界値分析"),
    ("F6","パスワード変更処理",
     "PW 変更フォーム + POST /api/auth/change-password",
     "新 PW / 確認 PW の組み合わせ × セッション種別",
     "変更成功 / フィールドエラー / API 拒否",
     "同値分割・状態遷移"),
    ("F7","セッション管理",
     "セッション Cookie + /api/auth/status, /api/auth/logout",
     "セッション状態（認証済み / ログアウト後 / 未認証）",
     "API のレスポンス / 画面のクリア状態",
     "状態遷移"),
    ("F8","ユーザー管理操作（管理者）",
     "POST /api/users, PUT /api/users/:id, DELETE /api/users/:id",
     "操作種別 × ユーザー現在状態",
     "ステータス変化 / 一覧への反映",
     "状態遷移・同値分割"),
]

PRIORITY_COLOR = {
    "最高": "C00000",
    "高":   "FF0000",
    "中":   "FF9900",
    "低":   "70AD47",
}

FRAME_COLOR = {
    "F1":"DEEBF7","F2":"E2EFDA","F3":"E2EFDA",
    "F4":"FFF2CC","F5":"FFF2CC","F6":"FCE4D6",
    "F7":"FCE4D6","F8":"EDE7F6","—":"F2F2F2",
}


# ─────────────────────────────────────────
# Excel 生成
# ─────────────────────────────────────────

def hex_fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def thin_border():
    s = Side(style="thin", color="BFBFBF")
    return Border(left=s, right=s, top=s, bottom=s)

def header_font():
    return Font(name="メイリオ", bold=True, color="FFFFFF", size=10)

def body_font(bold=False):
    return Font(name="メイリオ", bold=bold, size=9)

def wrap_align(h="left", v="top"):
    ha = {"left": "left", "center": "center", "right": "right"}[h]
    va = {"top": "top", "center": "center"}[v]
    return Alignment(horizontal=ha, vertical=va, wrap_text=True)

def set_col_width(ws, col_letter, cm_val):
    ws.column_dimensions[col_letter].width = cm_val * 4.72 / 1.8

def apply_header(ws, row, cols, bg="1F4E79"):
    for col in cols:
        c = ws.cell(row=row, column=col)
        c.fill = hex_fill(bg)
        c.font = header_font()
        c.alignment = wrap_align("center", "center")
        c.border = thin_border()

def write_cell(ws, row, col, value, bold=False, bg=None, align_h="left"):
    c = ws.cell(row=row, column=col, value=value)
    c.font = body_font(bold)
    c.alignment = wrap_align(align_h, "top")
    c.border = thin_border()
    if bg:
        c.fill = hex_fill(bg)
    return c

def generate_excel():
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # ── シート1: 観点ツリー ──────────────────────
    ws1 = wb.create_sheet("観点ツリー")
    ws1.sheet_view.showGridLines = False
    ws1.freeze_panes = "A3"

    ws1.merge_cells("A1:D1")
    title = ws1["A1"]
    title.value = "VSTeP テスト観点ツリー — 社内Webシステム ログイン・ユーザー管理機能"
    title.font = Font(name="メイリオ", bold=True, size=14, color="FFFFFF")
    title.fill = hex_fill("1F4E79")
    title.alignment = wrap_align("center", "center")
    ws1.row_dimensions[1].height = 28

    headers = ["大分類", "中分類", "小分類", "具体的な観点・確認ポイント"]
    for i, h in enumerate(headers, 1):
        c = ws1.cell(row=2, column=i, value=h)
        c.fill = hex_fill("2E75B6")
        c.font = header_font()
        c.alignment = wrap_align("center", "center")
        c.border = thin_border()
    ws1.row_dimensions[2].height = 20

    tree_data = [
        # (大分類, 中分類, 小分類, 観点)
        ("機能性","認証機能","正常認証","有効 ID × 正しい PW × 有効ユーザー → ログイン成功"),
        ("機能性","認証機能","正常認証","ホーム画面にユーザー名が表示される"),
        ("機能性","認証機能","認証失敗","存在しない ID → 認証エラー"),
        ("機能性","認証機能","認証失敗","PW 不一致 → 残り試行回数を表示"),
        ("機能性","認証機能","認証失敗","無効化ユーザー → 認証拒否"),
        ("機能性","認証機能","アカウントロック","1〜4回目の失敗 → ロックせず残り回数を表示"),
        ("機能性","認証機能","アカウントロック","5回目失敗（境界値）→ ロック発動"),
        ("機能性","認証機能","アカウントロック","ロック中：正しい PW でもログイン拒否・残り時間表示"),
        ("機能性","認証機能","アカウントロック","ロック解除：30分後自動解除 / 管理者手動解除"),
        ("機能性","入力バリデーション","ユーザーID","必須チェック：未入力 → エラーメッセージ"),
        ("機能性","入力バリデーション","ユーザーID","文字種チェック：半角英数字以外は無効"),
        ("機能性","入力バリデーション","ユーザーID","文字数：下限境界 0文字（無効）/ 1文字（有効）"),
        ("機能性","入力バリデーション","ユーザーID","文字数：上限境界 20文字（有効）/ 21文字（無効）"),
        ("機能性","入力バリデーション","パスワード","必須チェック：未入力 → エラーメッセージ"),
        ("機能性","入力バリデーション","パスワード","文字種チェック：スペース・全角・日本語は無効"),
        ("機能性","入力バリデーション","パスワード","文字数：下限境界 7文字（無効）/ 8文字（有効）"),
        ("機能性","入力バリデーション","パスワード","文字数：上限境界 32文字（有効）/ 33文字（無効）"),
        ("機能性","入力バリデーション","タイミング","フォーカスアウト（blur）時に該当フィールドのみ検証"),
        ("機能性","入力バリデーション","タイミング","ログインボタン押下時に全フィールドを検証"),
        ("機能性","パスワード期限管理","有効期限チェック","89日前（境界値-1）→ 期限切れでない"),
        ("機能性","パスワード期限管理","有効期限チェック","90日前（境界値）→ 期限切れ"),
        ("機能性","パスワード期限管理","有効期限チェック","91日前（境界値+1）→ 期限切れ"),
        ("機能性","パスワード期限管理","強制変更フロー","期限切れユーザーがログイン → PW 変更画面へ遷移"),
        ("機能性","パスワード期限管理","強制変更フロー","通常ログインユーザーは PW 変更 API を使えない"),
        ("機能性","パスワード変更処理","一致チェック","新 PW = 確認 PW → 変更成功"),
        ("機能性","パスワード変更処理","一致チェック","新 PW ≠ 確認 PW → フィールドエラー"),
        ("機能性","パスワード変更処理","変更後 UI","成功メッセージ表示・変更ボタンが disabled のまま"),
        ("機能性","パスワード変更処理","変更後 UI","2秒後に自動でログイン画面へ遷移"),
        ("機能性","セッション管理","ログイン後","/api/auth/status が 200 を返す"),
        ("機能性","セッション管理","ログアウト","ログイン画面に遷移・フォーム入力欄がクリアされる"),
        ("機能性","セッション管理","ログアウト","セッションが無効化される（ログアウト後の API が 401）"),
        ("機能性","ユーザー管理","ユーザー追加","正常追加（未使用 ID × 有効 PW）→ 一覧に反映"),
        ("機能性","ユーザー管理","ユーザー追加","ID 重複エラー → 400 DUPLICATE_USER"),
        ("機能性","ユーザー管理","ロック操作","手動ロック → badge-locked"),
        ("機能性","ユーザー管理","ロック操作","ロック解除 → badge-active・失敗カウントリセット"),
        ("機能性","ユーザー管理","有効/無効化","無効化 → badge-inactive（ログイン不可）"),
        ("機能性","ユーザー管理","有効/無効化","有効化 → badge-active（ログイン可）"),
        ("機能性","ユーザー管理","ユーザー削除","削除後に一覧から消える"),
        ("セキュリティ","インジェクション耐性","SQL インジェクション","ユーザーID フィールド：記号バリデーションで拒否"),
        ("セキュリティ","インジェクション耐性","SQL インジェクション","パスワードフィールド：認証失敗で終わり DB 改ざんなし"),
        ("セキュリティ","認証強度","ブルートフォース対策","5回失敗でロック発動"),
        ("セキュリティ","認証強度","パスワード保護","bcrypt ハッシュで保存（平文保存なし）"),
        ("セキュリティ","APIアクセス制御","未認証","未認証では保護 API が 401 を返す"),
        ("セキュリティ","APIアクセス制御","ログアウト後","ログアウト後の保護 API アクセスを 401 で拒否"),
        ("ユーザビリティ","エラーメッセージ","フィールドエラー","要件どおりのエラー文言・表示位置・色"),
        ("ユーザビリティ","エラーメッセージ","認証エラー","残り試行回数を明示"),
        ("ユーザビリティ","エラーメッセージ","ロックエラー","ロック状態・残り時間を明示"),
        ("ユーザビリティ","インタラクション","エラー消去","入力開始時に認証エラーが非表示になる"),
        ("信頼性","境界値保証","PW 文字数","7文字（無効）/ 8文字（有効）/ 32文字（有効）/ 33文字（無効）"),
        ("信頼性","境界値保証","ID 文字数","1文字（有効）/ 20文字（有効）/ 21文字（無効）"),
        ("信頼性","境界値保証","失敗回数","4回（ロックなし）/ 5回（ロック）"),
        ("信頼性","境界値保証","PW 期限","89日前（有効）/ 90日前（期限切れ）/ 91日前（期限切れ）"),
        ("信頼性","データ整合性","カウントリセット","ロック解除後に失敗カウントがリセットされる"),
        ("信頼性","データ整合性","日時更新","PW 変更後に password_changed_at が更新される"),
    ]

    cat_colors = {
        "機能性":   "DEEBF7",
        "セキュリティ": "FFF2CC",
        "ユーザビリティ":"E2EFDA",
        "信頼性":   "FCE4D6",
    }

    prev_l1 = prev_l2 = prev_l3 = None
    row = 3
    for l1, l2, l3, obs in tree_data:
        bg = cat_colors.get(l1, "FFFFFF")
        write_cell(ws1, row, 1, l1 if l1 != prev_l1 else "", bold=True, bg=bg, align_h="center")
        write_cell(ws1, row, 2, l2 if l2 != prev_l2 else "", bold=True, bg=bg)
        write_cell(ws1, row, 3, l3 if l3 != prev_l3 else "", bg=bg)
        write_cell(ws1, row, 4, obs, bg="FFFFFF")
        ws1.row_dimensions[row].height = 20
        prev_l1, prev_l2, prev_l3 = l1, l2, l3
        row += 1

    set_col_width(ws1, "A", 3.5)
    set_col_width(ws1, "B", 5.0)
    set_col_width(ws1, "C", 5.0)
    set_col_width(ws1, "D", 14.0)

    # ── シート2: テストフレーム ──────────────────
    ws2 = wb.create_sheet("テストフレーム")
    ws2.sheet_view.showGridLines = False
    ws2.freeze_panes = "A3"

    ws2.merge_cells("A1:F1")
    t2 = ws2["A1"]
    t2.value = "VSTeP テストフレーム — 社内Webシステム ログイン・ユーザー管理機能"
    t2.font = Font(name="メイリオ", bold=True, size=14, color="FFFFFF")
    t2.fill = hex_fill("1F4E79")
    t2.alignment = wrap_align("center", "center")
    ws2.row_dimensions[1].height = 28

    f_headers = ["フレームID","フレーム名","テスト対象","テスト条件（変数）","振る舞い（期待）","使用技法"]
    for i, h in enumerate(f_headers, 1):
        c = ws2.cell(row=2, column=i, value=h)
        c.fill = hex_fill("2E75B6")
        c.font = header_font()
        c.alignment = wrap_align("center", "center")
        c.border = thin_border()
    ws2.row_dimensions[2].height = 20

    frame_bg = ["DEEBF7","E2EFDA","E2EFDA","FFF2CC","FFF2CC","FCE4D6","FCE4D6","EDE7F6"]
    for i, (fid, fname, target, cond, behave, tech) in enumerate(FRAMES):
        r = i + 3
        bg = frame_bg[i]
        write_cell(ws2, r, 1, fid, bold=True, bg=bg, align_h="center")
        write_cell(ws2, r, 2, fname, bold=True, bg=bg)
        write_cell(ws2, r, 3, target, bg=bg)
        write_cell(ws2, r, 4, cond, bg=bg)
        write_cell(ws2, r, 5, behave, bg=bg)
        write_cell(ws2, r, 6, tech, bg=bg)
        ws2.row_dimensions[r].height = 36

    set_col_width(ws2, "A", 2.0)
    set_col_width(ws2, "B", 5.5)
    set_col_width(ws2, "C", 7.5)
    set_col_width(ws2, "D", 10.0)
    set_col_width(ws2, "E", 8.0)
    set_col_width(ws2, "F", 5.5)

    # ── シート3: テストコンディション一覧 ────────
    ws3 = wb.create_sheet("テストコンディション一覧")
    ws3.sheet_view.showGridLines = False
    ws3.freeze_panes = "A3"

    ws3.merge_cells("A1:H1")
    t3 = ws3["A1"]
    t3.value = "VSTeP テストコンディション一覧 — 社内Webシステム ログイン・ユーザー管理機能"
    t3.font = Font(name="メイリオ", bold=True, size=14, color="FFFFFF")
    t3.fill = hex_fill("1F4E79")
    t3.alignment = wrap_align("center", "center")
    ws3.row_dimensions[1].height = 28

    tc_headers = ["コンディションID","フレーム","テストグループ","テスト技法",
                  "テストコンディション（何を確認するか）",
                  "入力値 / 前提状態","期待結果","優先度"]
    for i, h in enumerate(tc_headers, 1):
        c = ws3.cell(row=2, column=i, value=h)
        c.fill = hex_fill("2E75B6")
        c.font = header_font()
        c.alignment = wrap_align("center", "center")
        c.border = thin_border()
    ws3.row_dimensions[2].height = 20

    for ri, tc in enumerate(TEST_CONDITIONS):
        r = ri + 3
        tc_id, frame, grp, tech, cond, inp, exp, prio = tc
        bg = FRAME_COLOR.get(frame, "FFFFFF")
        write_cell(ws3, r, 1, tc_id, bold=True, bg=bg, align_h="center")
        write_cell(ws3, r, 2, frame, bg=bg, align_h="center")
        write_cell(ws3, r, 3, grp, bg=bg)
        write_cell(ws3, r, 4, tech, bg=bg)
        write_cell(ws3, r, 5, cond, bg="FFFFFF")
        write_cell(ws3, r, 6, inp, bg="FFFFFF")
        write_cell(ws3, r, 7, exp, bg="FFFFFF")

        # 優先度セル：色付き
        pc = ws3.cell(row=r, column=8, value=prio)
        prio_color = PRIORITY_COLOR.get(prio, "000000")
        pc.font = Font(name="メイリオ", bold=True, color="FFFFFF", size=9)
        pc.fill = hex_fill(prio_color)
        pc.alignment = wrap_align("center", "center")
        pc.border = thin_border()

        ws3.row_dimensions[r].height = 30

    set_col_width(ws3, "A", 3.5)
    set_col_width(ws3, "B", 1.8)
    set_col_width(ws3, "C", 4.5)
    set_col_width(ws3, "D", 4.5)
    set_col_width(ws3, "E", 12.0)
    set_col_width(ws3, "F", 8.0)
    set_col_width(ws3, "G", 9.0)
    set_col_width(ws3, "H", 2.0)

    # オートフィルター
    ws3.auto_filter.ref = f"A2:H{len(TEST_CONDITIONS)+2}"

    # ── シート4: 思考プロセス（メモ） ────────────
    ws4 = wb.create_sheet("思考プロセス")
    ws4.sheet_view.showGridLines = False

    ws4.merge_cells("A1:B1")
    tp = ws4["A1"]
    tp.value = "VSTeP 設計思考プロセス — 観点ツリー・フレーム・コンディション作成の判断記録"
    tp.font = Font(name="メイリオ", bold=True, size=13, color="FFFFFF")
    tp.fill = hex_fill("1F4E79")
    tp.alignment = wrap_align("center", "center")
    ws4.row_dimensions[1].height = 26

    thought_rows = [
        ("フェーズ", "内容"),
        ("Step 1: 気になることの洗い出し（ボトムアップ）",
         "ユーザー視点・管理者視点・攻撃者視点・データ整合性視点の4視点から「気になること」を30項目ブレインストームした。\n"
         "例：「何回間違えるとロックされる？」「SQLインジェクションは防げる？」「期限切れ時に何が起きる？」"),
        ("Step 2: 抽象化・グルーピング",
         "30項目を「何の関心か」でまとめた。\n"
         "→ 認証処理 / PW期限管理 / 入力バリデーション / 画面UI / ユーザー管理 / セキュリティ / データ整合性の7カテゴリ\n"
         "→ さらに「機能性 / セキュリティ / ユーザビリティ / 信頼性」の4大分類に統合"),
        ("Step 3: トップダウンで分解・ツリー化",
         "各カテゴリを「なぜ？」「その中に何があるか？」と問いながら下に展開した。\n"
         "境界値（5回 / 90日 / 8文字 / 32文字）が出てきたら「閾値前 / ちょうど / 超過」の3点を必ず観点として追加した。"),
        ("フレーム切り方の判断基準",
         "「この観点では何が変数か？」を問い、変数1つに対してフレーム1つを作成した。\n"
         "例：失敗回数（1〜5）→ F4（ロック制御フレーム）\n"
         "　　経過日数（89/90/91）→ F5（PW期限フレーム）"),
        ("テスト技法の選択根拠",
         "・フレーム F1（認証）：ユーザー状態が離散的 → 同値分割 + 状態遷移\n"
         "・フレーム F2/F3（バリデーション）：上下限がある → 境界値分析 + 同値分割\n"
         "・フレーム F4（ロック）：5回という閾値 → 境界値分析（4回/5回）\n"
         "・フレーム F5（期限）：90日という閾値 → 境界値分析（89/90/91日）\n"
         "・フレーム F6〜F8：セッション状態・操作種別が主軸 → 状態遷移 + 同値分割"),
    ]

    for i, (ph, content) in enumerate(thought_rows):
        r = i + 2
        is_header = i == 0
        c1 = ws4.cell(row=r, column=1, value=ph)
        c2 = ws4.cell(row=r, column=2, value=content)
        if is_header:
            for c in [c1, c2]:
                c.fill = hex_fill("2E75B6")
                c.font = header_font()
                c.alignment = wrap_align("center", "center")
                c.border = thin_border()
            ws4.row_dimensions[r].height = 20
        else:
            c1.fill = hex_fill("D6E4F0")
            c1.font = Font(name="メイリオ", bold=True, size=9)
            c1.alignment = wrap_align("left", "top")
            c1.border = thin_border()
            c2.font = body_font()
            c2.alignment = wrap_align("left", "top")
            c2.border = thin_border()
            ws4.row_dimensions[r].height = 70

    ws4.column_dimensions["A"].width = 30
    ws4.column_dimensions["B"].width = 90

    path = OUT_DIR / "vstep-test-design.xlsx"
    wb.save(str(path))
    print(f"Excel 生成完了 → {path}")


# ─────────────────────────────────────────
# Word 生成
# ─────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_table_style(table, style="Table Grid"):
    table.style = style

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        run.font.size = Pt(13)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run.font.size = Pt(11)

def add_simple_table(doc, headers, rows, col_widths, header_bg="1F4E79", stripe=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
        cell.width = Cm(w)

    stripe_colors = ["FFFFFF", "EEF3FA"]
    for ri, row_data in enumerate(rows):
        bg = stripe_colors[ri % 2] if stripe else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ci == 0:
                cell.paragraphs[0].runs[0].font.bold = True
            set_cell_bg(cell, bg)
            cell.width = Cm(col_widths[ci])
    return table

def generate_word():
    doc = Document()

    # 余白設定
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # タイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VSTeP テスト設計書")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("社内Webシステム ログイン・ユーザー管理機能")
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

    # 基本情報表
    add_simple_table(doc,
        ["項目","内容"],
        [
            ("手法","VSTeP（Viewpoint-based Software Test design Process / 西 康晴 氏考案）"),
            ("対象システム","社内Webシステム ログイン・ユーザー管理機能（login-app）"),
            ("作成日","2026-05-14"),
            ("作成者","UedaKazuki"),
        ],
        [4.0, 13.0], header_bg="2E75B6"
    )
    doc.add_paragraph()

    # ─── 1. VSTePとは ───
    heading(doc, "1. VSTeP とは（手法概要）", 1)
    doc.add_paragraph(
        "VSTeP は「テスト観点」を中心に据えたテスト開発方法論です。\n"
        "テストとして気になること（観点）を洗い出し → ツリーに整理 → フレームで組み合わせ → "
        "コンディションで具体化する流れで進みます。"
    ).style = "Normal"

    add_simple_table(doc,
        ["フェーズ","成果物","問い"],
        [
            ("テスト要求分析","テスト観点ツリー","何をテストすべきか？"),
            ("テスト設計","テストフレーム","どの観点を組み合わせてテストするか？"),
            ("テスト詳細設計","テストコンディション一覧","具体的にどんな値・状態でテストするか？"),
        ],
        [4.0, 5.0, 8.0], header_bg="2E75B6"
    )
    doc.add_paragraph()

    # ─── 2. 思考プロセス：観点ツリー ───
    heading(doc, "2. 思考プロセス — 観点ツリーの構築", 1)
    doc.add_paragraph(
        "この節は「どのように観点ツリーを作ったか」の思考過程を記録したものです。"
    ).italic = True

    heading(doc, "Step 1 : 気になることのブレインストーム（ボトムアップ）", 2)
    doc.add_paragraph(
        "ログインシステムに対し、一般ユーザー・管理者・攻撃者・データ整合性の4視点から"
        "「気になること」を30項目出しました。"
    )
    concerns = [
        "【一般ユーザー】正しい ID/PW でログインできるか？間違ったPWを入れたらどうなるか？",
        "【一般ユーザー】何回間違えるとロックされるか？ロックはいつ自動解除されるか？",
        "【一般ユーザー】パスワードに有効期限があるか？期限が切れたとき何が起きるか？",
        "【一般ユーザー】入力欄に使えない文字（全角・記号・スペース）は何か？",
        "【一般ユーザー】空欄のまま送信したらどうなるか？PW が短すぎ・長すぎの場合は？",
        "【一般ユーザー】ログアウトしたらフォームはクリアされるか？",
        "【管理者】新しいユーザーを追加できるか？同じ ID で重複登録したらどうなるか？",
        "【管理者】管理者がアカウントをロック / 解除できるか？解除後に失敗カウントはリセットされるか？",
        "【管理者】ユーザーを無効化できるか？無効化後のログインは弾かれるか？",
        "【攻撃者】ユーザーID / パスワードへの SQL インジェクションを試みたら？",
        "【攻撃者】ログアウト後にセッション Cookie で API を叩けるか？",
        "【データ整合性】PW 変更後に password_changed_at は更新されるか？",
    ]
    for c in concerns:
        p = doc.add_paragraph(c, style="List Bullet")
        p.runs[0].font.size = Pt(9)

    heading(doc, "Step 2 : 抽象化・グルーピング", 2)
    doc.add_paragraph(
        "30項目を「何の関心か」でまとめ、7カテゴリ → 4大分類に整理しました。"
    )
    add_simple_table(doc,
        ["抽象カテゴリ","含まれる関心事"],
        [
            ("認証処理","正常ログイン / 認証失敗 / ロック"),
            ("PW期限管理","PW有効期限チェック / 強制変更フロー"),
            ("入力バリデーション","文字種 / 文字数（境界値）/ 必須チェック"),
            ("画面・UI動作","フォームクリア / 2秒後遷移 / ボタン状態"),
            ("ユーザー管理","追加 / ロック / 有効化 / 削除"),
            ("セキュリティ","SQLインジェクション / セッション管理 / 暗号化"),
            ("データ整合性","失敗カウントリセット / password_changed_at 更新"),
        ],
        [5.0, 12.0], header_bg="2E75B6"
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "さらに上位分類：機能性 / セキュリティ / ユーザビリティ / 信頼性 の4大分類に統合しました。"
    )

    heading(doc, "Step 3 : トップダウンで分解・ツリー化", 2)
    doc.add_paragraph(
        "各カテゴリを「なぜ？」「その中に何があるか？」と問いながら下に展開しました。\n"
        "境界値（5回 / 90日 / 8文字 / 32文字）が出てきたら、"
        "必ず「閾値前・ちょうど・超過」の3点を観点として追加しました。"
    )
    doc.add_paragraph()

    # ─── 3. 観点ツリー ───
    heading(doc, "3. テスト観点ツリー（最終版）", 1)
    tree_text = (
        "社内Webシステム ログイン・ユーザー管理機能\n"
        "│\n"
        "├── 【機能性】\n"
        "│   ├── 認証機能\n"
        "│   │   ├── 正常認証：有効 ID × 正しい PW × 有効ユーザー → ログイン成功\n"
        "│   │   ├── 認証失敗：存在しない ID / PW 不一致 / 無効化ユーザー\n"
        "│   │   └── アカウントロック\n"
        "│   │       ├── 1〜4回目 → 残り回数表示（ロックなし）\n"
        "│   │       ├── 5回目（境界値）→ ロック発動\n"
        "│   │       └── ロック中：正しい PW でも拒否・残り時間表示\n"
        "│   ├── 入力バリデーション\n"
        "│   │   ├── ユーザーID：必須 / 文字種（半角英数字）/ 文字数（1〜20文字）\n"
        "│   │   │   └── 境界値：0/1文字（下限）・20/21文字（上限）\n"
        "│   │   └── パスワード：必須 / 文字種（半角英数記号）/ 文字数（8〜32文字）\n"
        "│   │       └── 境界値：7/8文字（下限）・32/33文字（上限）\n"
        "│   ├── パスワード期限管理\n"
        "│   │   ├── 境界値：89日前（有効）/ 90日前（期限切れ）/ 91日前（期限切れ）\n"
        "│   │   └── 期限切れ → PW 変更画面へ強制遷移\n"
        "│   ├── パスワード変更処理\n"
        "│   │   ├── 新 PW = 確認 PW → 変更成功\n"
        "│   │   ├── 新 PW ≠ 確認 PW → フィールドエラー\n"
        "│   │   └── 変更成功後：ボタン disabled / 2秒後にログイン画面遷移\n"
        "│   ├── セッション管理\n"
        "│   │   ├── ログイン後：/api/auth/status が 200\n"
        "│   │   └── ログアウト：画面遷移・フォームクリア・セッション無効化\n"
        "│   └── ユーザー管理（管理者）\n"
        "│       ├── 追加（正常 / ID重複エラー）\n"
        "│       ├── ロック操作（手動ロック / 解除・失敗カウントリセット）\n"
        "│       ├── 有効/無効化\n"
        "│       └── ユーザー削除\n"
        "│\n"
        "├── 【セキュリティ】\n"
        "│   ├── SQL インジェクション耐性（ID フィールド / PW フィールド）\n"
        "│   ├── 認証強度：bcrypt ハッシュ保存 / 5回ロック\n"
        "│   └── セッション：ログアウト後の API 拒否（401）\n"
        "│\n"
        "├── 【ユーザビリティ】\n"
        "│   ├── フィールドエラー：文言・表示位置・色が仕様どおり\n"
        "│   ├── 認証エラー：残り試行回数を明示\n"
        "│   └── インタラクション：入力開始時にエラーが消える\n"
        "│\n"
        "└── 【信頼性・データ整合性】\n"
        "    ├── 境界値保証：PW/ID 文字数・失敗回数・PW 期限\n"
        "    └── 操作後データ更新：失敗カウントリセット・password_changed_at 更新"
    )
    p = doc.add_paragraph()
    run = p.add_run(tree_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    doc.add_paragraph()

    # ─── 4. 思考プロセス：テストフレーム ───
    heading(doc, "4. 思考プロセス — テストフレームの構築", 1)
    doc.add_paragraph(
        "テストフレームは「テスト対象 × テスト条件（変数）→ 振る舞い」の組み合わせ構造です。\n"
        "「この観点では何が変数か？」を問い、変数1つに対してフレーム1つを作成しました。"
    )

    add_simple_table(doc,
        ["判断の問い","対応フレーム"],
        [
            ("ユーザー状態（有効/ロック/無効/期限切れ）が変わったら？","F1 ログイン認証処理"),
            ("ユーザーID の入力値（空/文字種/文字数）が変わったら？","F2 ユーザーID バリデーション"),
            ("パスワードの入力値（空/文字種/文字数）が変わったら？","F3 パスワード バリデーション"),
            ("連続失敗回数（1〜5回）が変わったら？","F4 アカウントロック制御"),
            ("PW 変更日からの経過日数が変わったら？","F5 パスワード期限管理"),
            ("新 PW と確認 PW の組み合わせが変わったら？","F6 パスワード変更処理"),
            ("セッション状態（認証済み/未認証/ログアウト後）が変わったら？","F7 セッション管理"),
            ("管理者の操作種別 × ユーザー状態が変わったら？","F8 ユーザー管理操作"),
        ],
        [9.0, 8.0], header_bg="2E75B6"
    )
    doc.add_paragraph()

    # ─── 5. テストフレーム ───
    heading(doc, "5. テストフレーム", 1)

    frame_details = [
        ("F1 ログイン認証処理",
         [("テスト対象","POST /api/auth/login"),
          ("テスト条件","ユーザー状態（有効/無効/ロック/期限切れ）× ID・PW の正否"),
          ("振る舞い","HTTP ステータス・レスポンス内容・画面遷移先"),
          ("使用技法","同値分割・状態遷移")],
         ["#","ユーザー状態","入力 ID","入力 PW","期待 HTTP","期待 画面/メッセージ"],
         [
             ("1-1","有効","正しい","正しい","200","ホーム画面へ遷移"),
             ("1-2","有効","正しい","誤り（1回目）","401","「あと4回失敗するとロック」"),
             ("1-3","有効","正しい","誤り（4回目）","401","「あと1回失敗するとロック」"),
             ("1-4","有効","正しい","誤り（5回目）","403","ロックメッセージ"),
             ("1-5","ロック中","正しい","正しい","403","ロック中・残り時間"),
             ("1-6","無効","正しい","正しい","401","認証エラー"),
             ("1-7","存在しない","—","—","401","認証エラー"),
             ("1-8","期限切れ","正しい","正しい","422","PW変更画面へ遷移"),
         ],
         [0.8, 2.2, 2.0, 3.5, 2.0, 6.0]),
        ("F4 アカウントロック制御",
         [("テスト対象","連続失敗カウント機能 / ロック発動ロジック"),
          ("テスト条件","累積失敗回数（1〜5回）"),
          ("振る舞い","残り試行回数の表示 / ロック発動の有無"),
          ("使用技法","境界値分析・状態遷移")],
         ["#","累積失敗回数","残り回数","ロック発動","期待するメッセージ"],
         [
             ("4-1","1回目","4","なし","あと4回失敗するとロックされます"),
             ("4-2","2回目","3","なし","あと3回失敗するとロックされます"),
             ("4-3","3回目","2","なし","あと2回失敗するとロックされます"),
             ("4-4","4回目（境界値-1）","1","なし","あと1回失敗するとロックされます"),
             ("4-5","5回目（境界値）","0","発動","ロックメッセージ（auth-locked クラス）"),
             ("4-6","ロック後再試行","—","維持","アカウントがロックされています。約○分後に解除"),
         ],
         [0.8, 4.0, 2.0, 2.0, 7.0]),
        ("F5 パスワード期限管理",
         [("テスト対象","isPasswordExpired() 関数 / ログイン時の期限チェック"),
          ("テスト条件","password_changed_at からの経過日数"),
          ("振る舞い","正常ログイン成功 or PW 変更画面への遷移"),
          ("使用技法","境界値分析")],
         ["#","経過日数","閾値との関係","期待する結果"],
         [
             ("5-1","0日（今日）","閾値より大幅に前","期限切れでない → 正常ログイン"),
             ("5-2","89日前","閾値の1日前（境界値-1）","期限切れでない → 正常ログイン"),
             ("5-3","90日前","閾値ちょうど（境界値）","期限切れ → PW 変更画面へ遷移"),
             ("5-4","91日前","閾値の1日後（境界値+1）","期限切れ → PW 変更画面へ遷移"),
             ("5-5","180日前","大幅超過","期限切れ → PW 変更画面へ遷移"),
         ],
         [0.8, 2.5, 5.0, 8.0]),
    ]

    for frame_title, props, tc_headers, tc_rows, tc_widths in frame_details:
        heading(doc, frame_title, 2)
        add_simple_table(doc,
            ["要素","内容"],
            props, [3.0, 14.0], header_bg="2E75B6"
        )
        doc.add_paragraph()
        add_simple_table(doc, tc_headers, tc_rows, tc_widths, header_bg="595959")
        doc.add_paragraph()

    doc.add_paragraph(
        "※ フレーム F2/F3/F6/F7/F8 の詳細マトリクスは Excel ファイル（テストフレームシート）を参照してください。"
    ).italic = True
    doc.add_paragraph()

    # ─── 6. テストコンディション一覧 ───
    heading(doc, "6. テストコンディション一覧（全 41 件）", 1)

    heading(doc, "技法の選択根拠", 2)
    add_simple_table(doc,
        ["フレーム","選んだ技法","理由"],
        [
            ("F1（認証）","同値分割 + 状態遷移","ユーザー状態が離散的。同状態内の複数誤PW は代表1件で十分"),
            ("F2・F3（バリデーション）","境界値分析 + 同値分割","文字数に上下限 → 境界値必須。文字種は無効クラスから代表1件"),
            ("F4（ロック）","境界値分析","5回閾値に対して4回・5回が境界値"),
            ("F5（期限）","境界値分析","90日閾値に対して89/90/91日が境界値"),
            ("F6（PW変更）","同値分割 + 状態遷移","セッション状態と入力組み合わせ両方を確認"),
            ("F7（セッション）","状態遷移","認証状態は離散的な状態機械"),
            ("F8（管理）","状態遷移 + 同値分割","ユーザー状態遷移が主目的"),
        ],
        [2.5, 5.0, 9.0], header_bg="2E75B6"
    )
    doc.add_paragraph()

    heading(doc, "一覧表", 2)
    tc_table_headers = ["コンディションID","フレーム","テストグループ","テスト技法",
                        "テストコンディション","入力値 / 前提状態","期待結果","優先度"]
    tc_table_rows = [list(tc) for tc in TEST_CONDITIONS]
    tc_widths = [3.0, 1.5, 3.5, 3.5, 6.0, 4.5, 5.5, 1.5]

    table = doc.add_table(rows=1 + len(tc_table_rows), cols=len(tc_table_headers))
    table.style = "Table Grid"

    for i, (h, w) in enumerate(zip(tc_table_headers, tc_widths)):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1F4E79")
        cell.width = Cm(w)

    prio_hex = {"最高": "C00000", "高": "C55A11", "中": "BF8F00", "低": "375623"}
    stripe = ["FFFFFF", "EEF3FA"]

    for ri, row_data in enumerate(tc_table_rows):
        row_bg = stripe[ri % 2]
        for ci, val in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8)
            if ci == 0:
                run.font.bold = True
            if ci == 7:  # 優先度
                prio_val = str(val)
                hex_c = prio_hex.get(prio_val, "595959")
                set_cell_bg(cell, hex_c)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_cell_bg(cell, row_bg)
            cell.width = Cm(tc_widths[ci])

    doc.add_paragraph()

    # ─── 7. 付録 ───
    heading(doc, "7. 付録：作成ステップ整理", 1)
    steps = [
        ("Step 1","ブレインストーム（ボトムアップ）",
         "4視点（ユーザー/管理者/攻撃者/整合性）から「気になること」を30項目出した"),
        ("Step 2","抽象化・グルーピング",
         "7カテゴリ → 4大分類（機能性/セキュリティ/ユーザビリティ/信頼性）に整理"),
        ("Step 3","ツリー化（トップダウン）",
         "「なぜ？」「何が変わる？」と問いながら展開。境界値は「前・ちょうど・後」の3点を必ず追加"),
        ("Step 4","テストフレーム作成",
         "変数ごとにフレーム1つを作成。「テスト対象 × 条件 → 振る舞い」のマトリクスを記載"),
        ("Step 5","テストコンディション一覧",
         "各フレームのセルを1行1テストに具体化。技法の選択理由を明示しフレーム ID で追跡可能に"),
    ]
    add_simple_table(doc,
        ["ステップ","フェーズ","作業内容"],
        steps, [2.0, 5.0, 10.0], header_bg="2E75B6"
    )

    path = OUT_DIR / "vstep-test-design.docx"
    doc.save(str(path))
    print(f"Word 生成完了 → {path}")


if __name__ == "__main__":
    generate_excel()
    generate_word()
